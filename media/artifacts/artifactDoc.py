import json
import re
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

def generate_artifacts_doc(artifacts_json_path, output_path=None):
    # Load your JSON data
    with open(artifacts_json_path, 'r') as artifacts_file:
        artifacts_data = json.load(artifacts_file)

    # Get the directory where this script is located to find ArtifactInfo.json
    script_dir = os.path.dirname(os.path.abspath(__file__))
    info_file_path = os.path.join(script_dir, 'ArtifactInfo.json')
    
    with open(info_file_path, 'r') as info_file:
        artifact_info = json.load(info_file)

    # Create a new Document
    doc = Document()

    # Function to add spacing before heading and set font size
    def add_heading(doc, text, level=1):
        paragraph = doc.add_heading(text, level)
        run = paragraph.runs[0]
        run.font.size = Pt(14)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(12)

    # Function to get the interpreted field name from ArtifactInfo
    def get_interpreted_name(key, info_list):
        for item in info_list:
            if 'name' in item and key in item['name'].values():
                return list(item['name'].keys())[list(item['name'].values()).index(key)]
        return key

    # Function to find the set name from ArtifactInfo
    def find_set_name(set_key, tiers):
        for tier in tiers:
            for artifact_set in tier['data']['Sets']:
                if set_key in artifact_set['name'].values():
                    return list(artifact_set['name'].keys())[list(artifact_set['name'].values()).index(set_key)]
        return set_key

    # Function to find the substat name from ArtifactInfo
    def find_substat_name(substat_key, tiers):
        for tier in tiers:
            for substat in tier['data']['Substats']:
                if substat_key in substat['name'].values():
                    return list(substat['name'].keys())[list(substat['name'].values()).index(substat_key)]
        return substat_key

    # Function to convert rarity to star emojis
    def rarity_to_stars(rarity):
        return "⭐" * rarity

    # Function to split character name into two separate names
    def split_character_name(name):
        return " ".join(re.findall('[A-Z][^A-Z]*', name))

    # Iterate over the artifacts and add details to the document
    for index, artifact in enumerate(artifacts_data['artifacts']):
        set_name = find_set_name(artifact['setKey'], artifact_info['ArtifactTiers'])
        slot_name = get_interpreted_name(artifact['slotKey'], artifact_info['Pieces'])
        main_stat_name = get_interpreted_name(artifact['mainStatKey'], artifact_info['ArtifactTiers'][0]['data']['MainStats'])
        rarity_stars = rarity_to_stars(artifact['rarity'])
        character_name = split_character_name(artifact['location'])

        add_heading(doc, f"Artifact {index + 1}", level=1)
        
        # Create a table with 1 row and 2 columns
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Add artifact details to the first column
        cell = table.cell(0, 0)
        details_text = f"{set_name}\n{rarity_stars}\n{slot_name} | +{artifact['level']}\nMain Stat: {main_stat_name}"
        if artifact['location']:
            details_text += f"\nEquipped by {character_name}"
        else:
            details_text += "\n"  # Insert a blank line if the artifact is not equipped
        cell.text = details_text
        cell.paragraphs[0].runs[0].font.size = Pt(12)

        # Add substats to the second column, translated using ArtifactInfo
        cell = table.cell(0, 1)
        substats_text = "\n".join(
            [f"{find_substat_name(sub['key'], artifact_info['ArtifactTiers'])}: {sub['value']}" for sub in artifact['substats']]
        )
        cell.text = f"Sub Stats:"
        cv = round(sum(
            (
                sub['value'] * 2 if 'critRate' in sub['key'] else sub['value']
            ) for sub in artifact['substats'] if 'critRate' in sub['key'] or 'critDMG' in sub['key']
        ), 1)
        locked_status = ", Locked" if artifact.get('lock', False) else ""
        cell.text = f"Sub Stats: ({cv}cv{locked_status})"
        cell.text += f"\n{substats_text}"
        cell.paragraphs[0].runs[0].font.size = Pt(12)

    # Save the document
    # Generate the filename with the current date if no output path provided
    if output_path is None: output_path = f"artifacts.docx"

    # Save the document with the generated filename
    doc.save(output_path)
    return output_path

# For backwards compatibility when run as standalone script
if __name__ == '__main__':
    generate_artifacts_doc('artifacts.filtered.json', 'artifacts.docx')
    print("Document created successfully!")
